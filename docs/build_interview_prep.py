from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "docreview-agent-interview-prep.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(92, 99, 107)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_GOLD = "FFF8E8"
PALE_RED = "FCEBEC"
WHITE = RGBColor(255, 255, 255)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = tr_pr.find(qn("w:cantSplit"))
        if cant_split is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.keep_with_next = True


def set_font(run, name="Microsoft YaHei", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, name="Microsoft YaHei", size=11, color=None, bold=None):
    style.font.name = name
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = color
    if bold is not None:
        style.font.bold = bold


def set_cell_text(cell, text: str, *, bold=False, color=None, size=9.5) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.12
    run = p.add_run(text)
    set_font(run, size=size, color=color, bold=bold)


def add_table(doc: Document, headers: list[str], rows: Iterable[Iterable[str]], widths: list[int]) -> None:
    values = [list(row) for row in rows]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color=INK, size=9.5)
        set_cell_shading(table.rows[0].cells[idx], LIGHT_BLUE)
    repeat_table_header(table.rows[0])
    for row in values:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], str(value), size=9.2)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375 + level * 0.25)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_font(r, size=10.7)


def add_number(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    previous_is_number = len(doc.paragraphs) > 1 and doc.paragraphs[-2].style.name == "List Number"
    if previous_is_number:
        num_id = getattr(doc, "_interview_num_id")
    else:
        numbering = doc.part.numbering_part.element
        style_num_id = doc.styles["List Number"]._element.pPr.numPr.numId.val
        base_num = next(
            item for item in numbering.findall(qn("w:num"))
            if int(item.get(qn("w:numId"))) == style_num_id
        )
        abstract_id = base_num.find(qn("w:abstractNumId")).get(qn("w:val"))
        num_id = max(
            (int(item.get(qn("w:numId"))) for item in numbering.findall(qn("w:num"))),
            default=0,
        ) + 1
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract = OxmlElement("w:abstractNumId")
        abstract.set(qn("w:val"), abstract_id)
        num.append(abstract)
        level_override = OxmlElement("w:lvlOverride")
        level_override.set(qn("w:ilvl"), "0")
        start_override = OxmlElement("w:startOverride")
        start_override.set(qn("w:val"), "1")
        level_override.append(start_override)
        num.append(level_override)
        numbering.append(num)
        setattr(doc, "_interview_num_id", num_id)
    num_pr = p._p.get_or_add_pPr().get_or_add_numPr()
    num_pr.get_or_add_ilvl().val = 0
    num_pr.get_or_add_numId().val = num_id
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_font(r, size=10.7)


def add_para(doc: Document, text: str, *, bold_prefix: str | None = None, color=None, italic=False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_font(r1, size=10.8, bold=True, color=color)
        r2 = p.add_run(text[len(bold_prefix) :])
        set_font(r2, size=10.8, color=color, italic=italic)
    else:
        r = p.add_run(text)
        set_font(r, size=10.8, color=color, italic=italic)


def add_callout(doc: Document, label: str, text: str, fill: str = LIGHT_GRAY) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.18
    p.paragraph_format.left_indent = Pt(6)
    p.paragraph_format.right_indent = Pt(6)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "5")
        border.set(qn("w:color"), "7F8C99")
        borders.append(border)
    p_pr.append(borders)
    r1 = p.add_run(label + "  ")
    set_font(r1, size=10.5, bold=True, color=INK)
    r2 = p.add_run(text)
    set_font(r2, size=10.5)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_font(r, size={1: 16, 2: 13, 3: 12}[level], color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level], bold=True)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, size=11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for level, size, color, before, after in (
        (1, 16, BLUE, 18, 10),
        (2, 13, BLUE, 14, 7),
        (3, 12, DARK_BLUE, 10, 5),
    ):
        style = styles[f"Heading {level}"]
        set_style_font(style, size=size, color=color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15
    for name in ("List Bullet", "List Number"):
        style = styles[name]
        set_style_font(style, size=10.7)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_run = header.add_run("DocReview Agent | 面试准备参考")
    set_font(header_run, size=8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("Phase 0 文档 | 证据时间 2026-08-19")
    set_font(footer_run, size=8.5, color=MUTED)


def title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("AI Agent 面试准备文档")
    set_font(r, size=25, color=INK, bold=True)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(14)
    r = p2.add_run("DocReview Agent：从 HTTP 契约到持久化 Runtime 的系统化讲解与追问地图")
    set_font(r, size=14, color=MUTED)
    metadata = [
        ("项目定位", "Python 3.13 + FastAPI + PostgreSQL + LangGraph 的文档审查 AI Agent 服务"),
        ("文档用途", "面试前建立代码结构、业务目标、技术架构、核心数据流和追问应答的共同心智模型"),
        ("证据范围", "当前工作树实际 FastAPI 装配、src/docreview、tests、docs/remediation；不把历史目录当作在线能力"),
        ("阶段约束", "Phase 0：只读分析与治理文档；不连接数据库、不执行 migration、不调用真实 Provider/Tika"),
    ]
    add_table(doc, ["项目项", "结论"], metadata, [2100, 7260])
    add_callout(
        doc,
        "先记住一句话",
        "这是一个把模型推理变成可恢复、可审计、可授权、可重放的文档审查工作流。LangGraph 只做有界编排，PostgreSQL 才保存 Run/Step/Attempt/Tool/Approval/Commit/Outbox/Projection 事实。",
        LIGHT_BLUE,
    )


def build_content(doc: Document) -> None:
    title_block(doc)

    add_heading(doc, "1. 面试开场：30 秒和 2 分钟版本", 1)
    add_heading(doc, "30 秒版本", 2)
    add_para(doc, "DocReview Agent 是一个面向文档审查和修改的持久化 AI Agent 服务。外层用 FastAPI 暴露资源、文件、Assistant Turn、Run、Approval 和 SSE 接口；中间用 PostgreSQL 记录完整工作流事实和事务边界；LangGraph 只承载有限节点、checkpoint 和 interrupt/resume。典型链路是上传文档或提交 Turn，创建 Turn/Run/Step，Runtime worker 领取 Step，Graph 节点通过 RuntimeBoundary 调用模型、Context、Tool 和 Commit 子系统，结果写入 Outbox，再由 Projection worker 生成公开 DTO，HTTP 或 SSE 只读取 projection 并支持幂等重试和 Last-Event-ID 重放。")
    add_heading(doc, "2 分钟版本", 2)
    add_para(doc, "业务目标不是简单地让 LLM 输出一段文字，而是让审查结果可以回到准确的 Resource/Version/Node/Evidence，并在高风险修改前经过授权。文档先经过有界解析、规范化、AST/hash 和 chunk projection；检索侧区分 legacy Resource Search 与 EvidenceService，运行时还要固化 ContextManifest。Agent graph 先理解目标、组装上下文、决定下一动作，再检索/读节点/分析证据/生成 Patch；Patch 必须带 base version、expected node hash 和 evidence refs，经 Patch Validation 后创建绑定到 patch fact 的 Approval，外部 owner/admin 决定后才允许 Commit。Commit 在 Serializable 事务中重新检查 scope、版本和节点 hash，并和新版本、派生 projection、Outbox 一起原子提交。")
    add_callout(doc, "面试表达重点", "回答架构题时先讲事实源和不变量，再讲框架。不要把“用了 LangGraph”讲成系统设计本身。", PALE_GOLD)

    add_heading(doc, "2. 业务目标与在线范围", 1)
    add_heading(doc, "业务目标", 2)
    for item in [
        "让用户上传并阅读文档，在准确的资源和版本范围内检索证据。",
        "让 Agent 产生有 provenance 的观察、发现和 Patch，而不是只返回不可验证的自然语言。",
        "让修改动作可在审批、资源授权、版本冲突和幂等键约束下安全提交。",
        "让长任务在进程重启、SSE 断开、worker 抢占和 provider 重试后继续恢复，而不是依赖进程内状态。",
        "保持当前 HTTP、DTO、错误状态、SSE event name、X-Request-ID 和 Last-Event-ID 兼容性。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "实际注册的 API", 2)
    add_table(
        doc,
        ["路由族", "方法与路径", "关键契约"],
        [
            ("基础", "GET /healthz", "返回 {status: ok, service: server}; middleware 回写 X-Request-ID"),
            ("资源", "GET /api/resources; /{id}; /{id}/export; /{id}/search", "workspace scope、当前版本、最多 5 条 citation、导出 Markdown"),
            ("Agent Run", "GET /api/agent/runs; /{id}", "signed user identity；公开 allowlist DTO，不泄露 raw state/tool payload"),
            ("Approval", "GET /api/agent/approvals; /{id}; POST /{id}/approve|reject", "owner/admin 决策，状态冲突 409，decision 原子恢复 continuation"),
            ("Assistant 能力", "GET /api/assistant/capabilities", "上传扩展名由 parser policy 派生"),
            ("Assistant session", "GET /api/assistant/sessions; /{id}; DELETE /{id}", "当前使用 compatibility workspace scope"),
            ("Assistant Turn", "POST /api/assistant/conversations[ /stream ]; POST /sessions/{id}/messages[ /stream ]", "trusted ingress、durable acceptance、非流式 projection 或 SSE replay"),
            ("Assistant upload", "POST /api/assistant/conversations/files; /sessions/{id}/files", "multipart file，默认 20 MiB，上游 parser 和元数据事务"),
            ("文件", "GET /api/files/{id}/download", "按 workspace 查元数据，再流式读取原始内容"),
        ],
        [1600, 3900, 3860],
    )
    add_callout(doc, "范围陷阱", "当前应用未注册 resource task-context，历史 tasks/approvals/execution_jobs、shadow comparison 和 parity 包也不能自动算作在线闭包。", PALE_RED)

    add_heading(doc, "3. 代码结构：从入口读到核心事实", 1)
    add_table(
        doc,
        ["目录", "职责", "面试时的读法"],
        [
            ("api/", "FastAPI app factory、middleware、路由、依赖注入、错误映射、生产装配", "先看 main.py，再看 routes 与 api/assembly.py，确定真实在线范围"),
            ("identity/", "trusted ingress HMAC、WorkspaceScope、membership/role/resource policy", "回答租户隔离和为什么不能信任客户端 header"),
            ("turn/", "Turn DTO、幂等 acceptance、共享 HTTP/SSE pipeline、事件映射", "回答 request replay、非流式与流式一致性"),
            ("runtime/", "Run/Step/Attempt/Tool/Approval/Outbox runtime engine、lease、retry、projection worker", "回答重启恢复、并发 worker、fencing"),
            ("agent_graph/", "严格 Pydantic GraphState、节点、interrupt/resume、checkpointer、生产边界适配", "回答 LangGraph 到底负责什么"),
            ("tool_runtime/", "工具 Registry、schema、policy、approval、rate limit、audit、artifact 和 side effect recovery", "回答工具调用安全和幂等"),
            ("document/", "上传解析、canonical AST、稳定 node id/hash、Patch、Commit", "回答证据回溯和修改冲突"),
            ("knowledge/ + context/", "legacy search、EvidenceService、结构化 chunk、ContextManifest 组装", "回答 RAG 召回、扩窗和上下文预算"),
            ("providers/", "LLM、embedding、reranker、web search、Tika 和生产依赖装配", "回答 provider 超时、重试、配置 fail-closed"),
            ("storage/postgres/", "参数化 SQL、repository、事务、锁、scope predicate、outbox/receipt", "回答数据一致性；不要只看 Python 类名"),
            ("operations/ + deploy/", "reconciliation、重投影、容量告警、nginx trusted ingress", "回答上线、回滚和运维边界"),
        ],
        [1900, 3850, 3610],
    )
    add_heading(doc, "建议阅读顺序", 2)
    for item in [
        "AGENTS.md、docs/remediation/status.md、api-contract.md、persistence-contract.md：冻结边界和不变量。",
        "src/docreview/api/main.py：应用工厂、CORS、request ID、lifespan、router 注册。",
        "src/docreview/api/assembly.py：生产依赖闭包，确认哪些 repository/provider/worker 真正接上。",
        "src/docreview/turn/coordinator.py、pipeline.py、sse.py：理解请求接受、投影等待和 replay。",
        "src/docreview/runtime/engine.py、runtime_repository.py、runtime/projection.py：理解 worker 与事务。",
        "src/docreview/agent_graph/models.py、graph.py、runtime.py：理解 GraphState、节点和边界。",
        "src/docreview/agent_graph/production.py、tool_runtime/runtime.py：理解真实副作用如何被 policy/audit 包住。",
        "src/docreview/document/commit.py、storage/postgres/document_commit.py、knowledge/evidence_service.py：理解证据到提交。",
        "tests/：用失败路径和 SQL contract 验证上述叙述，而不是只用 README 猜。",
    ]:
        add_number(doc, item)

    add_heading(doc, "4. 技术架构与依赖装配", 1)
    add_heading(doc, "分层架构", 2)
    add_table(
        doc,
        ["层", "核心组件", "边界原则"],
        [
            ("HTTP/兼容层", "FastAPI routes、APIError、request ID、CORS", "保持公开方法/path/DTO/状态；只把公开 projection 转成响应"),
            ("身份与授权", "TrustedIngressAdapter、PolicyResolver、IdentityRepository", "HMAC attestation 证明 principal/workspace；membership 和 resource ownership 再授权"),
            ("持久化事实层", "TurnRepository、RuntimeRepository、CommitStore、UploadMetadataRepository", "PostgreSQL 事务、锁、unique key 和 scope predicate 是事实源"),
            ("编排层", "LangGraph StateGraph、checkpoint、interrupt/resume", "只做有界控制流；通过 RuntimeRequest/Response 调用权威子系统"),
            ("能力层", "ModelGateway、Evidence/Context、ToolRuntime、Committer", "每类副作用都有 typed contract、policy、audit 或事务边界"),
            ("异步 worker", "RuntimeWorker、ProjectionWorker", "claim + lease + generation fencing；Outbox publication 与事实写入解耦"),
            ("可观察性", "JSON logging、capacity alerts、reconciliation", "记录安全 metadata，不写正文、凭据或 provider body"),
        ],
        [1800, 3600, 3960],
    )
    add_heading(doc, "生产装配闭包", 2)
    add_para(doc, "生产 lifespan 的顺序很重要：先加载 provider dependencies，再打开 PostgreSQL pool，然后调用 assemble_production_repositories。装配过程创建 Resource/Run/Assistant/Identity/Turn/Runtime/Projection/Upload repository，构建 canonical committer、legacy search、DocumentUploadService、TurnCoordinator + DurableRunner；如果启用 runtime worker，则继续创建 ProjectRuntimeBoundary、LangGraph executor/checkpointer、RuntimeEngine、RuntimeWorker 和 ProjectionWorker。任一关键 provider、tokenizer、trusted ingress、worker id 或 database pool 缺失，production 都 fail closed。")
    add_table(
        doc,
        ["依赖", "来源", "缺失时的行为"],
        [
            ("Model/Embedding/Reranker/Tika/FileStore", "providers/assembly.py + Settings", "production lifespan 拒绝启动；开发模式可保持空依赖"),
            ("Trusted ingress", "Settings.trusted_ingress + TrustedIngressAdapter", "生产 repository assembly 抛错；持久化路由无签名返回 401"),
            ("DatabasePool", "create_database_pool(settings)", "未打开或 DATABASE_URL 缺失时不组装 production SQL repository"),
            ("Runtime worker", "RUNTIME_WORKER_ENABLED、RUNTIME_WORKER_ID", "未启用时可仅提供 HTTP/acceptance 依赖；启用但闭包不全则 fail closed"),
            ("Projection worker", "build_production_durable_runtime", "与 RuntimeLifecycle 一起启动；负责公开 outcome projection"),
        ],
        [2600, 3200, 3560],
    )

    add_heading(doc, "5. 核心数据流：面试必须能画出来", 1)
    add_heading(doc, "5.1 文档上传流", 2)
    add_number(doc, "Ingress 校验 trusted identity、workspace 和 multipart file；文件大小和扩展名在调用 writer 前检查。")
    add_number(doc, "DocumentUploadService 调 parser，生成规范文本/AST、content hash，并把 bytes stage 到 content-addressed FileStore。")
    add_number(doc, "UploadMetadataRepository 在一个事务中按 session -> resource/version -> uploaded_file -> assistant_message 顺序写入；session update 与文件发布完成后才 commit。")
    add_number(doc, "若 parser、file publication 或任一 SQL 写入失败，全部 metadata rollback；错误消息可以作为 assistant message 的事实保存。")
    add_number(doc, "后续 canonical commit 或 projection 由专门路径处理，不把 embedding/provider I/O 塞进 metadata acceptance 事务。")
    add_callout(doc, "常见追问", "为什么先 stage 文件、最后 promote？答案是让数据库 commit 前就能证明文件发布成功，避免出现 metadata 已提交但 blob 缺失的半个事实图。", LIGHT_BLUE)

    add_heading(doc, "5.2 Assistant Turn 非流式流", 2)
    add_number(doc, "HTTP middleware 选择或生成 X-Request-ID，并在响应中回写；持久化请求需要该准确值参与 HMAC canonical tuple 和 idempotency。")
    add_number(doc, "路由解析 message/resource_id，使用 TrustedIngressAdapter 验证 principal、organization、workspace、roles、issued_at 和 signature。")
    add_number(doc, "DurableRunner 调 TurnCoordinator.submit；coordinator 规范化 UUID、runtime_mode、principal scope，生成 canonical input JSON/hash。")
    add_number(doc, "TurnRepository 在一个 acceptance 事务中查或创建 session、user message、Turn、Run、初始 UnderstandGoal Step、有序事件和 Outbox。")
    add_number(doc, "runner 轮询公开 projection；只有事件已有确定 terminal/waiting status 且 projection sequence 已追平，才返回 DTO。超时返回 503，客户端用相同 request id 重试。")
    add_callout(doc, "关键思想", "HTTP 线程不是 Agent 生命周期。请求结束只代表 acceptance 或读取 projection，Run/Step/worker 仍可在进程外继续。", PALE_GOLD)

    add_heading(doc, "5.3 SSE 与 Last-Event-ID replay", 2)
    add_para(doc, "SSE 与非流式共用同一个 DurableRunner。stream 路由在启动前完成 body、cursor 和 identity validation；启动后 observer 把 TurnEvent 映射成冻结的 event name：turn_state、message_completed、error、done。持久化 sequence 直接作为 SSE id，重连时只 replay Last-Event-ID 之后的事件。传输取消只取消 observer，不取消持久化 pipeline task；客户端用同一 request_id 重连。")
    add_table(doc, ["状态/事实", "公开 SSE frame"], [
        ("turn.accepted / turn.running / run.queued", "turn_state"),
        ("assistant.message", "message_completed"),
        ("turn.waiting_input / turn.waiting_approval", "turn_state + done"),
        ("turn.succeeded", "done"),
        ("turn.failed / turn.cancelled", "error + done"),
        ("内部异常或 persisted terminal 无结果", "非负 reconnect cursor 的 error，随后关闭"),
    ], [3500, 5860])

    add_heading(doc, "5.4 Runtime worker 与 Graph 流", 2)
    add_number(doc, "RuntimeEngine.recover 先回收过期 lease，再 process_one；Repository 用 FOR UPDATE SKIP LOCKED claim 一个可运行 Step，并递增 lease_generation。")
    add_number(doc, "每次执行带 owner、lease expiry、generation 的 WorkItem；heartbeat、retry、completion 都带同样的 fencing predicate。")
    add_number(doc, "LangGraphExecutor 以 run_id + step namespace 建立 checkpoint。Graph node 本身不访问 provider/repository，而是 interrupt 发出 RuntimeRequest。")
    add_number(doc, "RuntimeBoundary 将 request 分派到 ModelGateway、ContextAssembler、ToolRuntime、Committer 或 Runtime；返回严格 RuntimeResponse 后 resume checkpoint。")
    add_number(doc, "遇到 await_user_input/await_approval，GraphExecutor 返回 WAIT_INPUT/WAIT_APPROVAL，并把 graph_request、checkpoint thread/step、graph_state 持久化到 Step output；外部事实到位后 continuation resume。")
    add_number(doc, "完成或失败时，RuntimeRepository 在事务中写 Attempt terminal telemetry、Step/Run outcome、next steps 和 Outbox。")

    add_heading(doc, "5.5 Agent graph 节点序列", 2)
    add_table(doc, ["节点", "输入/输出事实", "失败或停止条件"], [
        ("UnderstandGoal", "request fact -> Goal + context_manifest_id", "严格 schema；模型不能覆盖 scope"),
        ("AssembleContext", "Context candidates -> immutable ContextManifest", "token budget/reserved output 约束"),
        ("DecideNextAction", "Goal/context/observations/budget -> typed Action", "cycle、no-progress、budget exhaustion"),
        ("RetrieveEvidence / ReadDocumentNodes", "ToolRuntime observation -> persisted Observation", "policy、resource ownership、rate limit、tool error"),
        ("AnalyzeEvidence", "observations -> FindingRef + observation", "finding 必须绑定 evidence ids"),
        ("GeneratePatch", "finding refs -> PatchRef", "base version、operation shape、patch hash"),
        ("ValidatePatch", "PatchRef -> valid PatchRef", "AST/schema/hash/source authorization"),
        ("RequestApproval", "valid patch -> pending ApprovalRef", "审批 request 只能创建 pending fact"),
        ("AwaitApproval", "pending approval + checkpoint -> approved/rejected", "决定必须匹配 approval/patch fact"),
        ("CommitPatch", "approved patch -> CommitRef + Outbox", "Serializable recheck；失败即 conflict"),
        ("RenderOutcome", "facts -> OutcomeRef / public message", "只输出 bounded projection"),
    ], [2000, 4000, 3360])

    add_heading(doc, "5.6 高风险修改的完整闭环", 2)
    add_para(doc, "高风险工具在 ToolDefinition 中必须 requires_approval=True。RuntimeToolExecutor 的顺序是：严格 JSON/schema 校验 -> scope/resource binding -> Policy -> approval binding -> rate limit -> audit claim -> backend execute/recover -> output schema/size/artifact -> audit finish。Graph 不能绕过 ToolRuntime 直接写数据库。Approval fact 绑定 workspace、run、step、tool/version、input hash、patch hash、resource refs 和 target idempotency key；owner/admin 决策后，approval repository 原子地创建唯一 CommitPatch continuation 或 rejected terminal outcome。")

    add_heading(doc, "6. 一致性、安全和可恢复性不变量", 1)
    add_table(doc, ["不变量", "实现证据", "面试答法"], [
        ("幂等", "Turn/Run/Step/Tool/Approval/Commit/Outbox/Projection receipt unique keys + canonical hash", "相同 key 相同 input 返回既有事实；不同 canonical body 必须 conflict"),
        ("隔离", "SQL 每条 query/write 都绑定 workspace/resource/version；Policy 再查 membership/ownership", "不能先跨租户查再内存过滤"),
        ("fencing", "owner + lease_expires_at + lease_generation 条件", "旧 worker 即使恢复也不能覆盖新 claimant"),
        ("事务", "acceptance/outcome/approval/commit/upload 各自闭包", "Outbox 是事务交接点，provider I/O 在事务外"),
        ("公开面", "Run DTO allowlist、Projection reader、SSE mapper", "不能暴露 raw state、manifest、tool payload、凭据或内部 trace"),
        ("可恢复", "checkpoint + graph_resume + persisted events + projection receipts", "重启/断线后重读事实，不依赖内存 channel"),
        ("预算", "max_steps、max_tool_calls、token/cost/deadline + cycle/no-progress", "模型不能无限循环，也不能无限放大上下文"),
        ("审计", "Attempt、Tool audit、Observation、Approval、Commit、Outbox", "每个副作用有可解释的 provenance 和状态迁移"),
    ], [1750, 4050, 3560])
    add_heading(doc, "Trusted ingress 的签名 tuple", 2)
    add_para(doc, "HMAC-SHA256 canonical input 为 v1、request_id、HTTP_METHOD、request_path、principal_type、principal_id、organization_id、workspace_id、issued_at_rfc3339nano、comma_separated_roles 按换行连接。Adapter 检查 UUID、principal type、时间窗口、低写十六进制 signature、常量时间比较和请求 workspace 一致性。Ingress 必须先剥离客户端伪造 header，再生成 attestation；Python 适配器不是 IdP。")

    add_heading(doc, "7. 文档审查与检索设计", 1)
    add_heading(doc, "当前已实现的文档事实", 2)
    for item in [
        "Document model 保存 NodeType、source_location、page_mapping、metadata、content_hash；stable_node_id 和 hash_node 让 Patch 能引用准确节点。",
        "Canonical commit 会重新读取/校验 base version、node hash、Patch hash、evidence 和 scope，并写完整 resource version、node/source mapping、section/chunk projection、document_patch_commit 和 document.version.committed Outbox。",
        "Legacy Resource Search 通过 embedding/reranker adapter 生成最多五条 citation；EvidenceService 提供更丰富的 channel、fusion、rerank、provenance 和 EvidenceSet 类型。",
        "ContextManifest 是不可变事实，保存 tokenizer、token budget、reserved output、items、total_tokens 和 content_hash；replay 不能从当前检索结果重建。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "待批准的结构化父子切块规格，不是当前完成能力", 2)
    add_para(doc, "docs/remediation/document-review-chunking-spec.md 明确目标是：结构化 parser 保留 heading/list/table/page mapping；tokenizer-aware child chunks 用于 lexical/semantic recall；parent window 在 rerank 后按 window_group_id 有界扩展；每个扩展 child 保留独立 provenance；ContextManifest 固化后 replay 不再查 chunks。当前缺陷包括固定 800 字符、长 section 边界不足、Tika 只请求 text/plain、section_type/metadata 不完整、window_group_id 粗粒度和 golden 覆盖不足。面试时必须把“规格目标”和“当前生产现状”分开。")
    add_table(doc, ["阶段", "数据", "不能做的事"], [
        ("解析", "bounded parser -> ParsedElement / AST", "不能让 LLM 推断标题或结构事实"),
        ("规范化", "stable node id/hash + source/page mapping", "不能改变现有 hash 编码和 patch 语义"),
        ("投影", "section + child + parent window + embedding profile", "不能在 HTTP read 中 lazy backfill"),
        ("召回", "lexical/semantic/rerank child candidates", "不能跨 workspace/resource/version/profile 召回后内存过滤"),
        ("扩窗", "ordered sibling child -> ContextManifest items", "不能扩大 Patch evidence 授权范围"),
    ], [1700, 4050, 3610])

    add_heading(doc, "8. 面试官最可能问什么：按优先级准备", 1)
    add_heading(doc, "P0：一定会问的架构题", 2)
    p0 = [
        ("请你从一个 Assistant Turn 讲完整请求链路。", "从 request ID/trusted ingress 开始，讲 acceptance 事务、Run/Step、worker claim、Graph interrupt、Tool/Context/Commit、Outbox、Projection、HTTP/SSE；明确 HTTP 线程不拥有执行生命周期。", "如果 SSE 断了怎么办？如果同 request id 重试怎么办？"),
        ("为什么已经有 LangGraph，还要 Run/Step/Attempt/Tool 这些表？", "Graph state/checkpoint 是可重建编排状态，不能替代跨进程事实、审计、授权、lease 和公开 projection；数据库事实支持重启、并发和对账。", "checkpoint 丢了能否恢复？哪些状态能重新计算？"),
        ("如何防止两个 worker 同时执行同一个 Step？", "FOR UPDATE SKIP LOCKED claim + owner + expiry + generation；heartbeat/completion/retry 都重复 predicate；旧 claimant 更新 影响行数为 0 即 LeaseLost。", "为什么只靠 asyncio lock 不够？"),
        ("如何保证重试不产生重复副作用？", "稳定 idempotency key、canonical input hash、audit claim/recovery receipt；side-effecting tool 没有 backend receipt 时 fail closed，不盲目重放。", "相同 key 但 input 改了怎么办？"),
        ("为什么 Outbox 必须和事实写在同一事务？", "Outbox 是事实到异步投影/发布的原子交接；同事务保证不出现事实已提交但没有事件，或事件先于事实。发布本身可重试，receipt 保证幂等。", "为什么不直接在 commit 事务里调用 provider/webhook？"),
        ("workspace/resource/principal 隔离在哪里做？", "Ingress 证明 scope，handler 绑定 scope，Policy 查 membership/ownership，SQL 每个边界再次带 workspace/resource/version predicate，Commit 再 Serializable recheck。", "能否先查全量再 Python 过滤？"),
        ("审批如何避免只在内存里挂起？", "Approval 是持久事实，绑定 patch/input/resource/hash；Graph checkpoint 记录 await request；决策事务原子创建 continuation 或 rejection outcome，再由 worker resume。", "批准后如何证明批准的是原来那份 patch？"),
        ("Commit 为什么还要重新校验 hash 和版本？", "模型生成和审批之间文档可能变化；Serializable lock/recheck 防止 stale base version、node hash、scope 或 evidence 被绕过。", "如果 base version 冲突返回什么？"),
    ]
    for question, answer, follow in p0:
        add_heading(doc, question, 3)
        add_para(doc, "答题主线：" + answer)
        add_para(doc, "最可能追问：" + follow, color=MUTED)

    add_heading(doc, "P1：深入实现题", 2)
    p1 = [
        ("DurableRunner 为什么要等 projection 追平事件 sequence？", "避免返回一个 events 已 terminal 但 public DTO 还没写好的半成品；deterministic status + projection.status + last_event_sequence 三者同时满足才返回。"),
        ("Last-Event-ID 重放为什么不能只靠内存队列？", "内存队列会随进程消失，持久化 TurnEvent sequence 才能跨重启重放；SSE id 直接使用事实 sequence。"),
        ("Graph node 如何限制死循环？", "GraphLimits 的 max_cycles、max_no_progress、max_observations，加上 Runtime 的 max_steps、max_tool_calls、token/cost/deadline。重复 observation hash 增加 no_progress。"),
        ("ToolRuntime 的安全执行顺序是什么？", "schema -> scope/resource -> policy -> approval -> rate limit -> audit claim -> backend/recovery -> output schema/size/artifact -> audit finish。前置失败不得产生后续副作用。"),
        ("大工具输出为什么要 artifactize？", "inline output 有字节和 token 上限；大内容写 content-addressed artifact，只在公开观察中返回 bounded summary 和 provenance。"),
        ("为什么公开 Run DTO 不暴露 GraphState？", "GraphState 含内部 fact ids、manifest、tool input/output、checkpoint 和控制状态；公开 projection 只给用户需要的 Run/Step/ToolCall/Approval/Finding allowlist。"),
        ("上传事务为什么锁 session，且 message sequence 用数据库生成？", "串行化 session 内消息顺序，避免并发写产生相同 sequence；同一事务让 session、resource、file、message 一起成功或回滚。"),
        ("当前 API 认证覆盖是否完全对称？", "不是。Run/Approval、持久化 Turn 和 upload 校验 trusted ingress；health、resources、capabilities、session query/delete、file download 走 compatibility scope 或无 identity adapter，这是冻结兼容缺口，不能口头声称全站已租户安全。"),
        ("配置如何避免开发模式误判生产可用？", "development 可启动但 AppDependencies 为空；production lifespan 检查 provider、pool、trusted ingress、worker、tokenizer 并 fail closed。"),
        ("数据库 round-trip 测试为何可能没有跑？", "仓库有测试 fuse：只有 ALLOW_DB_TESTS=1、TEST_DATABASE_URL、_test 数据库名和 host allowlist 全满足才连接；否则必须 skip，不能读生产 DATABASE_URL。"),
    ]
    for question, answer in p1:
        add_heading(doc, question, 3)
        add_para(doc, "答题主线：" + answer)

    add_heading(doc, "P2：针对当前代码缺口和兼容性的追问", 2)
    p2 = [
        ("为什么缺 resource_id 的 durable Turn 不是 400？", "Handler DTO 允许可选，但当前 durable repository/prepare 要求非空 scope；冻结行为会落到通用 500，这是兼容性缺口，不能在未批准的修复中擅自改成 400。"),
        ("相同 request id 不同 body 为什么不是 409？", "内部有 IdempotencyConflictError，但当前 Assistant adapter 的冻结映射仍是通用 500 或已启动 SSE 后的 error event；改变公开状态属于单独 API 修复。"),
        ("为什么 docs 里写了父子切块，却不能说已上线？", "规格文件标明待批准实施，且 status 只保证现有 deterministic chunk/projection 接口；面试回答必须列出现状缺陷和实施门禁。"),
        ("仓库没有 migrations，如何证明 SQL 合法？", "以批准的 schema artifact 和现有 SQL contract tests 为依据；Phase 0 不根据 Python SQL 臆造 migration，也不连接数据库。"),
        ("为什么一些 read route 没有 trusted ingress？", "当前兼容性边界认证不对称；只能如实指出风险和后续修复，不把 Runtime scope 的存在当作所有路由已隔离。"),
        ("生产是否真的提供了 web search / Tika / file store？", "依赖由 providers assembly 和 settings 装配；当前环境没有真实外部验证 artifact，因此只能说接口和 fail-closed 装配存在，不能声称 staging 已通过。"),
    ]
    for question, answer in p2:
        add_heading(doc, question, 3)
        add_para(doc, "答题主线：" + answer)

    add_heading(doc, "9. 面试官从哪里开始追问：一条可预测的深挖路径", 1)
    add_table(doc, ["起点", "第一层追问", "第二层追问", "第三层追问"], [
        ("Assistant Turn", "如何接受请求？", "如何幂等和重放？", "projection lag / timeout / SSE cancel 怎么处理？"),
        ("LangGraph", "有哪些节点？", "interrupt/resume 如何实现？", "为什么 checkpoint 不是事实源？"),
        ("Tool call", "工具如何注册？", "policy/approval/rate limit 顺序？", "side effect timeout 后如何恢复而不重放？"),
        ("Document patch", "Patch 如何表达？", "expected_hash/base_version 如何校验？", "审批后文档变化如何避免 stale commit？"),
        ("PostgreSQL", "哪些表保存 Run？", "事务边界是什么？", "SKIP LOCKED + generation 如何 fencing？"),
        ("RAG", "如何搜证据？", "lexical/semantic/rerank 如何融合？", "父窗口扩展是否扩大授权？"),
        ("安全", "谁给 workspace？", "HMAC 如何防伪造？", "为什么部分兼容路由仍无 trusted identity？"),
        ("上线", "如何启动？", "怎么观察 queue/lease/outbox lag？", "回滚如何保留已接受事实并继续恢复？"),
    ], [1750, 2470, 2470, 2670])
    add_callout(doc, "答题顺序", "先说状态/事实，再说事务/锁，再说失败恢复，最后说 API 映射。面试官通常会从“功能怎么跑”追到“并发和失败时谁是事实源”。", LIGHT_BLUE)

    add_heading(doc, "10. 代码题/系统设计题的回答模板", 1)
    add_heading(doc, "模板 A：设计一个可恢复的 Agent worker", 2)
    for item in [
        "定义事实：Run、Step、Attempt、Tool、Approval、Outbox、Projection。",
        "定义状态机和合法迁移；把 terminal/waiting 状态写入数据库。",
        "用事务 + SKIP LOCKED claim；给 worker owner、expiry、generation。",
        "每次 heartbeat/retry/complete 都带 fencing predicate。",
        "副作用使用业务幂等 key 和 input hash；side effect receipt 不确定时 fail closed。",
        "Outbox 与事实同事务；Projection 以 event receipt 幂等。",
        "重启先 recover lease，再领取新任务；监控 queue age、lease、retry、dead letter。",
    ]:
        add_number(doc, item)
    add_heading(doc, "模板 B：排查“用户收到 200 但没有结果”", 2)
    for item in [
        "检查 X-Request-ID、Turn 是否 acceptance 成功，input hash 是否冲突。",
        "检查 TurnEvent 是否已有 deterministic terminal/waiting 状态。",
        "检查 Run/Step 是否被 worker claim，lease 是否过期，Attempt 是否 terminal。",
        "检查 Outbox 是否 pending/publishing/dead_letter，projection receipt 是否存在。",
        "检查 public projection status/last_event_sequence 是否追平。",
        "确认不要用同一个请求创建第二个 Turn；用同 request id 重试或 SSE Last-Event-ID replay。",
    ]:
        add_number(doc, item)
    add_heading(doc, "模板 C：解释一次 Patch 不应提交的原因", 2)
    add_para(doc, "先区分：schema invalid、policy denied、approval rejected、idempotency conflict、base version conflict、expected node hash mismatch、scope mismatch、lease lost。每一类都应在正确的事实层记录 error category 和审计，而不是把所有错误转成模型自然语言。Commit 事务必须 rollback，不应写半个新版本或孤立 Outbox。")

    add_heading(doc, "11. 当前已知缺口、面试时如何诚实表达", 1)
    add_table(doc, ["事实", "不要夸大的说法", "推荐说法"], [
        ("Phase 0 门禁", "已经做完全链路生产验证", "当前完成代码/契约和 database-free/static 验证；真实 DB/provider/Tika/staging 需受控环境"),
        ("数据库", "仓库自带完整 migrations", "仓库不含历史 migration 文件，SQL 行为以 schema owner artifact 和 contract test 为准"),
        ("结构化切块", "父子 window 已经全部上线", "已有 chunk/evidence/context 类型和待批准目标规格，规格中也记录当前固定字符 splitter 的缺陷"),
        ("认证", "所有 API 都有统一租户认证", "持久化 Agent/Turn/upload 有 trusted ingress；部分兼容读路由仍是非对称边界"),
        ("外部依赖", "provider/Tika 已在生产通过", "装配和 fail-closed 条件已定义，当前未在本环境执行真实外部流量验证"),
        ("测试", "所有测试都访问真实数据库", "默认数据库 fuse 阻止连接；SQL 形态、失败路径和静态装配可无数据库验证"),
    ], [1900, 3150, 4310])
    add_heading(doc, "12. 证据索引：被追问时打开哪些文件", 1)
    add_table(doc, ["主题", "首选文件", "可证明的结论"], [
        ("在线路由", "src/docreview/api/main.py; api/routes/*", "注册的 API、middleware、错误和路由边界"),
        ("生产装配", "src/docreview/api/assembly.py; providers/assembly.py; runtime/assembly.py", "依赖闭包、fail-closed、worker lifecycle"),
        ("Turn/SSE", "src/docreview/turn/coordinator.py; pipeline.py; sse.py", "canonical input/hash、projection wait、event mapping/replay"),
        ("Runtime", "src/docreview/runtime/engine.py; runtime_repository.py; runtime_sql.py", "claim、heartbeat、retry、outcome、outbox SQL"),
        ("Graph", "src/docreview/agent_graph/models.py; graph.py; runtime.py", "strict schema、节点、interrupt、checkpoint namespace"),
        ("Tool", "src/docreview/tool_runtime/runtime.py; models.py; registry.py", "schema/policy/approval/rate-limit/audit/artifact 顺序"),
        ("身份", "src/docreview/identity/trusted_ingress.py; policy.py", "HMAC tuple、时间窗、membership/resource policy"),
        ("文档/提交", "src/docreview/document/model.py; commit.py; storage/postgres/document_commit.py", "AST/hash、Patch 校验、Serializable commit、Outbox"),
        ("上传", "src/docreview/document/upload.py; storage/postgres/upload_write.py", "stage/promote 和 session/resource/file/message 事务顺序"),
        ("契约", "docs/remediation/api-contract.md; persistence-contract.md", "冻结 HTTP、SSE、事务、幂等、scope、lease 语义"),
        ("验证", "tests/api; tests/storage; tests/runtime; tests/agent_graph; tests/tool_runtime", "失败路径、SQL contract、公开 DTO 和 graph/tool 边界"),
    ], [1800, 4400, 3160])

    add_heading(doc, "13. 一页速记", 1)
    add_callout(doc, "业务", "文档 -> 证据 -> 发现 -> Patch -> 审批 -> Commit；结果必须可回溯、可授权、可恢复。", LIGHT_BLUE)
    add_callout(doc, "架构", "FastAPI 是兼容边界；PostgreSQL 是事实源；LangGraph 是有界控制流；Outbox/Projection 是公开结果桥。", LIGHT_BLUE)
    add_callout(doc, "并发", "SKIP LOCKED claim + owner/expiry/generation fencing；旧 worker 永远不能覆盖新 claimant。", LIGHT_BLUE)
    add_callout(doc, "幂等", "canonical JSON/hash + stable key；同 key 同 body replay，不同 body conflict。", LIGHT_BLUE)
    add_callout(doc, "安全", "trusted ingress 证明 principal/workspace，Policy 查 membership/ownership，SQL 每个边界再次 scope。", LIGHT_BLUE)
    add_callout(doc, "SSE", "事实 sequence 就是 SSE id；Last-Event-ID 只 replay 后续事件；断线不取消持久化执行。", LIGHT_BLUE)
    add_callout(doc, "诚实边界", "当前缺少真实 DB/provider/staging 证据；结构化父子切块是待批准规格；认证覆盖存在冻结的不对称。", PALE_GOLD)

    add_heading(doc, "附录：面试前自测清单", 1)
    for item in [
        "不用看代码，能在白板上画 Turn -> Run -> Step -> Attempt -> Tool/Approval -> Commit -> Outbox -> Projection。",
        "能说清为什么 GraphState、SSE queue、UI state 都不是事实源。",
        "能说出同 request id 相同/不同 body 的行为，以及当前 API 的兼容性例外。",
        "能解释 lease_generation 如何阻止旧 worker 写入，以及 recover 在启动时做什么。",
        "能按顺序讲 ToolRuntime 的 schema、scope、policy、approval、rate limit、audit 和 artifact。",
        "能说明 Patch 的 base_version、expected_hash、evidence_refs 和审批绑定。",
        "能区分当前 legacy search、已有 EvidenceService 类型、以及待批准的 parent-window 目标规格。",
        "能说出至少三条当前缺口，并给出不夸大的验证边界。",
        "能从 tests/storage 中指出一条 SQL contract，而不是只说“有测试”。",
        "能回答回滚不会删除已接受事实，而是切入口、排空 worker、继续使用同一 PostgreSQL facts。",
    ]:
        add_bullet(doc, item)


def main() -> None:
    doc = Document()
    configure_document(doc)
    build_content(doc)
    doc.core_properties.title = "DocReview Agent AI Agent 面试准备文档"
    doc.core_properties.subject = "项目代码结构、业务目标、技术架构、数据流与面试追问"
    doc.core_properties.author = "Codex"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
